import io
import json
import re
import os
from typing import Dict, Any

import streamlit as st
from google import genai
from pypdf import PdfReader
from docx import Document

st.set_page_config(
    page_title="Resume ATS Analyzer",
    page_icon="📄",
    layout="wide",
)

# -----------------------------
# Helpers
# -----------------------------
def get_secret(name: str) -> str:
    """Read a Streamlit secret safely."""
    try:
        value = st.secrets.get(name, "")
        if value:
            return str(value).strip()
    except Exception:
        pass
    return os.getenv(name, "").strip()


def extract_pdf_text(file_bytes: bytes) -> str:
    reader = PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        pages.append(page.extract_text() or "")
    return "\n".join(pages).strip()


def extract_docx_text(file_bytes: bytes) -> str:
    document = Document(io.BytesIO(file_bytes))
    parts = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            parts.append(paragraph.text.strip())

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts).strip()


def extract_resume_text(uploaded_file) -> str:
    data = uploaded_file.getvalue()
    suffix = uploaded_file.name.lower()

    if suffix.endswith(".pdf"):
        return extract_pdf_text(data)
    if suffix.endswith(".docx"):
        return extract_docx_text(data)

    raise ValueError("Unsupported file type. Please upload a PDF or DOCX file.")


def basic_ats_checks(text: str) -> Dict[str, Any]:
    """A deterministic baseline score before Gemini analysis."""
    normalized = re.sub(r"\s+", " ", text).strip()
    lower = normalized.lower()

    if not normalized:
        return {"score": 0, "checks": [], "word_count": 0}

    words = re.findall(r"\b[\w+#.-]+\b", normalized)
    word_count = len(words)

    checks = []

    section_patterns = {
        "Contact information": r"\b(email|e-mail|phone|mobile|linkedin)\b",
        "Summary/Profile": r"\b(summary|profile|objective|professional summary)\b",
        "Experience": r"\b(experience|employment|work history|professional experience)\b",
        "Education": r"\beducation\b",
        "Skills": r"\b(skills|technical skills|core competencies)\b",
    }

    for name, pattern in section_patterns.items():
        checks.append({
            "name": name,
            "passed": bool(re.search(pattern, lower)),
        })

    bullet_count = len(re.findall(r"(?:^|\n)\s*(?:[-•*▪◦]|\d+[.)])\s+", text))
    checks.append({"name": "Bullet-based achievements", "passed": bullet_count >= 3})

    action_words = (
        "achieved", "built", "created", "developed", "designed", "improved",
        "implemented", "increased", "reduced", "led", "managed", "optimized",
        "automated", "delivered", "analyzed", "engineered", "launched"
    )
    action_count = sum(len(re.findall(rf"\b{re.escape(w)}\b", lower)) for w in action_words)
    checks.append({"name": "Action-oriented language", "passed": action_count >= 3})

    quantified = len(re.findall(
        r"\b\d+(?:\.\d+)?\s*(?:%|percent|k|m|million|thousand|hours?|days?|months?|years?|users?|clients?)\b",
        lower,
    ))
    checks.append({"name": "Quantified achievements", "passed": quantified >= 2})

    # Simple formatting-risk signals visible in extracted text.
    checks.append({
        "name": "Reasonable resume length",
        "passed": 250 <= word_count <= 1200,
    })

    passed = sum(1 for c in checks if c["passed"])
    score = round(100 * passed / len(checks))

    return {
        "score": score,
        "checks": checks,
        "word_count": word_count,
    }


def parse_json_response(raw: str) -> Dict[str, Any]:
    """Extract JSON even if the model surrounds it with markdown fences."""
    cleaned = raw.strip()
    cleaned = re.sub(r"^```(?:json)?\s*", "", cleaned, flags=re.IGNORECASE)
    cleaned = re.sub(r"\s*```$", "", cleaned)

    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        match = re.search(r"\{.*\}", cleaned, flags=re.DOTALL)
        if match:
            return json.loads(match.group(0))
        raise


def analyze_with_gemini(resume_text: str, job_description: str, api_key: str) -> Dict[str, Any]:
    client = genai.Client(api_key=api_key)

    job_context = job_description.strip() or "No job description supplied. Evaluate general ATS readiness."

    prompt = f"""
You are an expert ATS resume evaluator and career coach.

Analyze the resume below. This is an ATS-readiness assessment, not a claim about
the score produced by any particular ATS vendor because ATS systems use different
scoring methods.

Return ONLY valid JSON. No markdown and no extra text.

Required JSON schema:
{{
  "ats_score": 0,
  "summary": "short overall assessment",
  "section_scores": {{
    "keyword_optimization": 0,
    "formatting_and_parseability": 0,
    "experience_and_achievements": 0,
    "skills": 0,
    "clarity_and_readability": 0
  }},
  "strengths": ["..."],
  "improvements": [
    {{
      "priority": "High|Medium|Low",
      "issue": "...",
      "recommendation": "..."
    }}
  ],
  "missing_keywords": ["..."],
  "keyword_matches": ["..."],
  "rewrite_examples": [
    {{
      "before": "...",
      "after": "..."
    }}
  ],
  "final_verdict": "..."
}}

Rules:
- ats_score must be an integer from 0 to 100.
- Be evidence-based. Do not invent experience, qualifications, metrics, employers,
  degrees, certifications, or skills.
- If a job description is supplied, assess keyword relevance against it.
- If no job description is supplied, assess general ATS readiness.
- Prefer measurable, specific improvements.
- Flag tables, columns, graphics, headers/footers, icons, unusual symbols, or
  other formatting risks only when the extracted text gives evidence or the user
  should verify them visually.
- Do not recommend keyword stuffing. Recommend natural keyword usage.
- Keep the response useful but concise.

JOB DESCRIPTION:
{job_context}

RESUME:
{resume_text[:30000]}
"""

    response = client.models.generate_content(
        model="gemini-3.6-flash",
        contents=prompt,
    )

    if not response.text:
        raise RuntimeError("Gemini returned an empty response.")

    return parse_json_response(response.text)


def combine_scores(baseline_score: int, ai_score: int) -> int:
    # AI gets more weight because it evaluates context and wording.
    return max(0, min(100, round((baseline_score * 0.30) + (ai_score * 0.70))))


# -----------------------------
# UI
# -----------------------------
st.title("📄 Resume ATS Analyzer")
st.write(
    "Upload a resume to get an ATS-readiness score, strengths, missing keywords, "
    "and practical improvements powered by Gemini Flash."
)

with st.sidebar:
    st.header("Settings")
    st.caption("Model: Gemini 3.6 Flash")
    st.info(
        "ATS scores are estimates. Different applicant-tracking systems use "
        "different parsing and ranking rules."
    )

api_key = get_secret("GEMINI_API_KEY")

if not api_key:
    st.warning(
        "GEMINI_API_KEY is not configured. Add it to Streamlit Secrets before "
        "running an AI analysis."
    )

uploaded_file = st.file_uploader(
    "Upload your resume",
    type=["pdf", "docx"],
    help="PDF or DOCX. For best ATS analysis, use a text-based resume rather than a scanned image.",
)

job_description = st.text_area(
    "Optional: paste the job description",
    height=220,
    placeholder="Paste the target job description here for job-specific keyword matching...",
)

analyze_button = st.button(
    "🔍 Analyze Resume",
    type="primary",
    use_container_width=True,
)

if analyze_button:
    if uploaded_file is None:
        st.error("Please upload a PDF or DOCX resume first.")
        st.stop()

    if not api_key:
        st.error("GEMINI_API_KEY is missing. Configure it in Streamlit Secrets.")
        st.stop()

    try:
        with st.spinner("Extracting resume text..."):
            resume_text = extract_resume_text(uploaded_file)

        if len(resume_text.strip()) < 80:
            st.error(
                "Very little text could be extracted. If this is a scanned/image-only "
                "PDF, convert it to a text-based PDF or DOCX and try again."
            )
            st.stop()

        baseline = basic_ats_checks(resume_text)

        with st.spinner("Gemini is analyzing your resume..."):
            ai_result = analyze_with_gemini(
                resume_text=resume_text,
                job_description=job_description,
                api_key=api_key,
            )

        ai_score = int(ai_result.get("ats_score", 0))
        final_score = combine_scores(baseline["score"], ai_score)

        st.session_state["analysis"] = {
            "file_name": uploaded_file.name,
            "baseline": baseline,
            "ai": ai_result,
            "final_score": final_score,
        }

    except Exception as exc:
        st.error(f"Analysis failed: {exc}")
        st.info(
            "Check that the Gemini API key is valid, the resume contains selectable "
            "text, and the file is a valid PDF/DOCX."
        )

analysis = st.session_state.get("analysis")

if analysis:
    st.divider()
    st.subheader(f"Results for {analysis['file_name']}")

    score = analysis["final_score"]
    c1, c2, c3 = st.columns(3)
    c1.metric("ATS Readiness Score", f"{score}/100")
    c2.metric("AI Score", f"{analysis['ai'].get('ats_score', 0)}/100")
    c3.metric("Extracted Words", analysis["baseline"]["word_count"])

    st.progress(score / 100)

    ai = analysis["ai"]

    st.markdown("### 🧾 Overall assessment")
    st.write(ai.get("summary", "No summary returned."))

    st.markdown("### 📊 Category scores")
    section_scores = ai.get("section_scores", {})
    if section_scores:
        cols = st.columns(len(section_scores))
        for col, (name, value) in zip(cols, section_scores.items()):
            pretty_name = name.replace("_", " ").title()
            col.metric(pretty_name, f"{value}/100")

    left, right = st.columns(2)

    with left:
        st.markdown("### ✅ Strengths")
        strengths = ai.get("strengths", [])
        if strengths:
            for item in strengths:
                st.write(f"• {item}")
        else:
            st.write("No strengths returned.")

    with right:
        st.markdown("### 🔑 Keyword analysis")
        matched = ai.get("keyword_matches", [])
        missing = ai.get("missing_keywords", [])

        st.write("**Matched / relevant:**")
        st.write(", ".join(matched) if matched else "None identified.")

        st.write("**Potentially missing:**")
        st.write(", ".join(missing) if missing else "None identified.")

    st.markdown("### 🛠️ Improvements")
    improvements = ai.get("improvements", [])
    if improvements:
        for item in improvements:
            priority = item.get("priority", "Medium")
            issue = item.get("issue", "")
            recommendation = item.get("recommendation", "")
            with st.expander(f"{priority} priority — {issue}"):
                st.write(recommendation)
    else:
        st.write("No improvement items returned.")

    st.markdown("### ✍️ Rewrite examples")
    examples = ai.get("rewrite_examples", [])
    if examples:
        for example in examples:
            st.markdown(f"**Before:** {example.get('before', '')}")
            st.markdown(f"**After:** {example.get('after', '')}")
            st.divider()
    else:
        st.write("No rewrite examples returned.")

    st.markdown("### 🎯 Final verdict")
    st.write(ai.get("final_verdict", "No final verdict returned."))

    with st.expander("Show extracted resume text"):
        st.text(resume_text if "resume_text" in locals() else "")
